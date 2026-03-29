"""
retrieval.py — Document parsing and BM25 retrieval.

Reads the admissions DOCX template file, splits it into sections
by detecting bold uppercase headings, and indexes them with BM25.
"""

import re
import copy
from docx import Document
from rank_bm25 import BM25Okapi


# ── Heading Detection ──

def is_template_heading(paragraph):
    """
    Detect template headings by checking:
    1. Text is uppercase and at least 10 chars
    2. Runs are bold (True) or inherited bold (None)
    3. No run is explicitly NOT bold (False)
    """
    text = (paragraph.text or "").strip()
    if not text or len(text) < 10:
        return False

    runs = paragraph.runs
    if not runs:
        return False

    if text != text.upper():
        return False

    has_explicit_bold = any(r.bold is True for r in runs)
    all_inherited = all(r.bold is None for r in runs)
    has_not_bold = any(r.bold is False for r in runs)

    return (has_explicit_bold or all_inherited) and not has_not_bold


# ── Section Parser ──

def build_sections(doc):
    """Parse the DOCX into template sections using bold+uppercase headings."""
    paragraphs = doc.paragraphs
    heading_indices = [i for i, p in enumerate(paragraphs) if is_template_heading(p)]

    raw_sections = []
    for idx, h_idx in enumerate(heading_indices):
        end_idx = (
            heading_indices[idx + 1] - 1
            if idx + 1 < len(heading_indices)
            else len(paragraphs) - 1
        )
        raw_sections.append({
            "title": paragraphs[h_idx].text.strip(),
            "start": h_idx,
            "end": end_idx,
        })

    skip_titles = {"ADMISSIONS FREQUENTLY ASKED INQUIRY QUESTIONS TEMPLATES"}
    cleaned = []
    i = 0

    while i < len(raw_sections):
        s = raw_sections[i]

        if s["title"] in skip_titles:
            i += 1
            continue

        if s["title"] == "OFFICE USE ONLY":
            if i + 1 < len(raw_sections):
                raw_sections[i + 1]["start"] = s["start"]
            i += 1
            continue

        if (i + 1 < len(raw_sections)
                and raw_sections[i + 1]["title"].startswith("EXEMPTIONS")):
            s["title"] = s["title"] + " " + raw_sections[i + 1]["title"]
            s["end"] = raw_sections[i + 1]["end"]
            i += 2
        else:
            i += 1

        para_indices = list(range(s["start"], s["end"] + 1))
        text_lines = []
        for j in para_indices:
            t = (paragraphs[j].text or "").strip()
            if t:
                text_lines.append(t)

        body_text = "\n".join(text_lines[1:]) if len(text_lines) > 1 else ""

        cleaned.append({
            "title": s["title"],
            "start": s["start"],
            "end": s["end"],
            "para_indices": para_indices,
            "text": body_text,
        })

    return cleaned


# ── Tokenizer ──

def tokenize(text):
    return re.findall(r"[a-zA-Z0-9']+", text.lower())


# ── Template Retriever ──

class TemplateRetriever:
    """Loads a DOCX, parses sections, builds BM25 index."""

    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.sections = build_sections(self.doc)

        # Build BM25 index with title boosted 3x
        corpus = [
            tokenize(s["title"] + " " + s["title"] + " " + s["title"] + " " + s["text"])
            for s in self.sections
        ]
        self.bm25 = BM25Okapi(corpus)

    def retrieve(self, query, k=5):
        """Return top-k matching sections with scores."""
        scores = self.bm25.get_scores(tokenize(query))
        ranked = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:k]
        return [{**self.sections[i], "score": float(scores[i])} for i in ranked]

    def fill_placeholders(self, section, student_info):
        """
        Fill placeholders in a deep copy of the doc.
        Returns the modified doc and section reference.
        """
        working_doc = copy.deepcopy(self.doc)
        name = student_info.get("name", "")
        semester = student_info.get("semester", "")

        for idx in section["para_indices"]:
            para = working_doc.paragraphs[idx]
            for run in para.runs:
                if not run.text:
                    continue

                # Semester
                if semester:
                    for variant in ["[specific semester]", "[Specific Semester]", "[SPECIFIC SEMESTER]"]:
                        if variant in run.text:
                            run.text = run.text.replace(variant, semester)
                else:
                    for variant in ["[specific semester]", "[Specific Semester]", "[SPECIFIC SEMESTER]"]:
                        if variant in run.text:
                            run.text = run.text.replace(variant, "the upcoming semester")

                # Name
                if name:
                    for greeting in [
                        "Dear Applicant,", "Dear Applicant",
                        "Dear Applican,", "Dear Applican",
                        "Dear Student,", "Dear Student",
                    ]:
                        if greeting in run.text:
                            replacement = (greeting
                                .replace("Applicant", name)
                                .replace("Applican", name)
                                .replace("Student", name))
                            run.text = run.text.replace(greeting, replacement)

        return working_doc

    def export_docx(self, doc, section, output_path):
        """Export a section as a standalone DOCX with formatting preserved."""
        new_doc = Document()
        for p in new_doc.paragraphs:
            p._element.getparent().remove(p._element)

        body = new_doc._element.body
        for idx in section["para_indices"]:
            if idx == section["start"]:
                continue
            body.append(copy.deepcopy(doc.paragraphs[idx]._p))

        new_doc.save(output_path)
        return output_path

    def get_plain_text(self, doc, section):
        """Get plain text of a section from a (possibly modified) doc."""
        lines = []
        for idx in section["para_indices"]:
            if idx == section["start"]:
                continue
            text = doc.paragraphs[idx].text.strip()
            if text:
                lines.append(text)
        return "\n".join(lines)

    def get_template_list(self):
        """Return list of all template titles."""
        return [s["title"] for s in self.sections]
