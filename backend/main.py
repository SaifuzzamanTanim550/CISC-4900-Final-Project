"""
BC Admissions Email Assistant — Backend
All logic in one file matching the Colab notebook.
Run with: uvicorn main:app --host 0.0.0.0 --reload
"""

import os
import re
import copy
import json
import uuid
from datetime import datetime, timezone
import requests as http_requests
from pathlib import Path

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from docx import Document
from docx.oxml.ns import qn
from rank_bm25 import BM25Okapi

# ── Load .env ──
load_dotenv()
HF_TOKEN = os.getenv("HF_TOKEN", "")
if not HF_TOKEN:
    print("WARNING: HF_TOKEN not set in .env")

MONGO_URI = os.getenv("MONGO_URI", "")

MODEL_ID = "meta-llama/Llama-3.1-8B-Instruct"
ROUTER_URL = "https://router.huggingface.co/v1/chat/completions"
CONFIDENCE_THRESHOLD = 8.0

# ── MongoDB setup ──
db = None
try:
    if MONGO_URI:
        from pymongo import MongoClient
        mongo_client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=5000)
        db = mongo_client["bc_admissions"]
        mongo_client.admin.command('ping')
        print("MongoDB connected.")
    else:
        print("WARNING: MONGO_URI not set — logging disabled.")
except Exception as e:
    print(f"WARNING: MongoDB connection failed — logging disabled. Error: {e}")
    db = None

# ── Load DOCX template ──
TEMPLATE_DIR = Path(__file__).parent / "templates"
DOCX_FILES = list(TEMPLATE_DIR.glob("*.docx"))
if not DOCX_FILES:
    raise FileNotFoundError(f"No .docx in {TEMPLATE_DIR}. Put your template there.")

DOCX_PATH = str(DOCX_FILES[0])
ORIGINAL_DOC = Document(DOCX_PATH)
print(f"Loaded: {DOCX_PATH} ({len(ORIGINAL_DOC.paragraphs)} paragraphs)")

# ── Output dir ──
OUTPUT_DIR = Path(__file__).parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)


# ═══════════════════════════════════════════════════════════════
# MONGODB LOGGING — saves anonymized query data
# ═══════════════════════════════════════════════════════════════

def log_query(data):
    """Log a query to MongoDB. Fails silently if DB is not connected."""
    if db is None:
        return None
    try:
        result = db.queries.insert_one({
            "timestamp": datetime.now(timezone.utc),
            "topic": data.get("topic", ""),
            "template_title": data.get("template_title", ""),
            "confidence": data.get("confidence", 0),
            "matched": data.get("matched", False),
            "feedback": None,
        })
        return str(result.inserted_id)
    except Exception as e:
        print(f"MongoDB log error: {e}")
        return None


# ═══════════════════════════════════════════════════════════════
# LLM HELPER
# ═══════════════════════════════════════════════════════════════

def llm_call(system_prompt, user_prompt, max_tokens=300, temperature=0):
    r = http_requests.post(
        ROUTER_URL,
        headers={
            "Authorization": f"Bearer {HF_TOKEN}",
            "Content-Type": "application/json",
        },
        json={
            "model": MODEL_ID,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "max_tokens": max_tokens,
            "temperature": temperature,
        },
        timeout=60,
    )
    if r.status_code != 200:
        raise RuntimeError(r.text)
    return r.json()["choices"][0]["message"]["content"].strip()


# ═══════════════════════════════════════════════════════════════
# HEADING PARSER
# ═══════════════════════════════════════════════════════════════

def is_template_heading(paragraph):
    text = (paragraph.text or "").strip()
    if not text or len(text) < 10:
        return False
    runs = paragraph.runs
    if not runs:
        return False
    if text != text.upper():
        return False
    has_bold = any(r.bold is True for r in runs)
    return has_bold


def build_sections(doc):
    paragraphs = doc.paragraphs
    heading_indices = [i for i, p in enumerate(paragraphs) if is_template_heading(p)]
    sections = []
    for idx, h_idx in enumerate(heading_indices):
        end_idx = heading_indices[idx + 1] - 1 if idx + 1 < len(heading_indices) else len(paragraphs) - 1
        para_indices = list(range(h_idx, end_idx + 1))
        text_lines = []
        for j in para_indices:
            t = (paragraphs[j].text or "").strip()
            if t:
                text_lines.append(t)
        body_text = "\n".join(text_lines[1:]) if len(text_lines) > 1 else ""
        sections.append({
            "title": paragraphs[h_idx].text.strip(),
            "start": h_idx,
            "end": end_idx,
            "para_indices": para_indices,
            "text": body_text,
        })
    return sections


def tokenize(text):
    return re.findall(r"[a-zA-Z0-9']+", text.lower())


# ── Build BM25 index (title boosted 3x) ──
SECTIONS = build_sections(ORIGINAL_DOC)
corpus = [
    tokenize(s["title"] + " " + s["title"] + " " + s["title"] + " " + s["text"])
    for s in SECTIONS
]
BM25 = BM25Okapi(corpus)
print(f"Templates: {len(SECTIONS)}")


# ═══════════════════════════════════════════════════════════════
# EXTRACT STUDENT INFO
# ═══════════════════════════════════════════════════════════════

def extract_student_info(email_text):
    system = """extract info from student email.
only return json.
fields:
name
semester
topic

do not guess name.
if no name, return empty string."""

    response = llm_call(system, email_text, max_tokens=150)
    try:
        match = re.search(r'\{[^}]+\}', response)
        info = json.loads(match.group()) if match else {}
    except Exception:
        info = {}

    name = info.get("name", "")
    fake_names = ["john doe", "jane doe", "applicant", "student", "n/a", "none", "unknown"]
    if name.lower().strip() in fake_names:
        name = ""

    return {
        "name": name,
        "semester": info.get("semester", ""),
        "topic": info.get("topic", "general inquiry"),
    }


# ═══════════════════════════════════════════════════════════════
# TEMPLATE SELECTION
# ═══════════════════════════════════════════════════════════════

def retrieve_top_k(email, k=5):
    scores = BM25.get_scores(tokenize(email))
    ranked = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:k]
    return [{**SECTIONS[i], "score": float(scores[i])} for i in ranked]


def llm_choose(email, candidates):
    num = len(candidates)
    previews = []
    for i, c in enumerate(candidates, 1):
        body_preview = "\n".join(c["text"].splitlines()[:10])
        previews.append(f"--- TEMPLATE {i} ---\n{c['title']}\n{body_preview}")

    templates_block = "\n\n".join(previews)

    system = f"""you help match student emails to response templates at Brooklyn College admissions.

instructions:
1. read the student email carefully
2. identify what they need help with
3. read each template below
4. if a template directly helps with their need, pick it
5. if no template helps, pick 0

{templates_block}

respond in this exact format only:
NEED: <what the student is asking about>
CHOICE: <number 1-{num} or 0 if no template helps with their need>
CONFIDENCE: <0 to 100, how well the template answers their need>
REASON: <why you picked this template or why none fit>

important:
- a template must actually address the student's need, not just share similar words
- if the student has a login or technical problem and no template covers that, pick 0
- if the student asks about a topic and a template covers that topic, pick it
- confidence above 80 means the template clearly answers their question
- confidence below 60 means you are not sure, pick 0 instead"""

    response = llm_call(system, email, max_tokens=120, temperature=0)

    choice = -1
    confidence = 0
    reason = ""

    for line in response.splitlines():
        line = line.strip()
        if line.upper().startswith("CHOICE:"):
            val = line.split(":", 1)[1].strip()
            for char in val:
                if char.isdigit():
                    c = int(char)
                    choice = c - 1 if c > 0 else -1
                    break
        elif line.upper().startswith("CONFIDENCE:"):
            val = line.split(":", 1)[1].strip()
            nums = re.findall(r"\d+", val)
            if nums:
                confidence = min(int(nums[0]), 100)
        elif line.upper().startswith("REASON:"):
            reason = line.split(":", 1)[1].strip()

    if confidence < 60:
        return {"choice": -1, "confidence": confidence, "reason": reason}

    return {"choice": choice, "confidence": confidence, "reason": reason}


LLM_CONFIDENCE_THRESHOLD = 60

def choose_template(email):
    candidates = retrieve_top_k(email, k=10)
    if candidates[0]["score"] < CONFIDENCE_THRESHOLD:
        return None
    result = llm_choose(email, candidates)
    if result["choice"] == -1 or result["confidence"] < LLM_CONFIDENCE_THRESHOLD:
        return None
    chosen = candidates[result["choice"]]
    chosen["llm_confidence"] = result["confidence"]
    chosen["llm_reason"] = result["reason"]
    return chosen


# ═══════════════════════════════════════════════════════════════
# FILL PLACEHOLDERS
# ═══════════════════════════════════════════════════════════════

def fill_placeholders(doc, section, student_info):
    name = student_info.get("name", "")
    semester = student_info.get("semester", "")

    for idx in section["para_indices"]:
        para = doc.paragraphs[idx]
        for run in para.runs:
            if not run.text:
                continue
            if semester:
                run.text = run.text.replace("[specific semester]", semester)
            else:
                run.text = run.text.replace("[specific semester]", "the upcoming semester")
            if name:
                run.text = run.text.replace("Dear Applicant", f"Dear {name}")
                run.text = run.text.replace("Dear Student", f"Dear {name}")


def export_doc(doc, section, out="response.docx"):
    new_doc = Document()
    for p in new_doc.paragraphs:
        p._element.getparent().remove(p._element)
    body = new_doc._element.body
    for idx in section["para_indices"]:
        if idx == section["start"]:
            continue
        body.append(copy.deepcopy(doc.paragraphs[idx]._p))
    new_doc.save(out)
    return out


def get_plain_text(doc, section):
    lines = []
    for idx in section["para_indices"]:
        if idx == section["start"]:
            continue
        text = doc.paragraphs[idx].text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════
# HTML CONVERTER
# ═══════════════════════════════════════════════════════════════

def _is_run_bold(run):
    if run.bold is True:
        return True
    if run.style and run.style.font and run.style.font.bold is True:
        return True
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        b = rPr.find(qn('w:b'))
        if b is not None:
            val = b.get(qn('w:val'))
            return val is None or val != '0'
    return False


def _is_run_italic(run):
    if run.italic is True:
        return True
    if run.style and run.style.font and run.style.font.italic is True:
        return True
    return False


def _is_run_underline(run):
    if run.underline is True:
        return True
    if run.style and run.style.font and run.style.font.underline is True:
        return True
    return False


def get_html_text(doc, section):
    html_parts = []
    in_list = False
    last_was_empty = False

    for idx in section["para_indices"]:
        if idx == section["start"]:
            continue
        para = doc.paragraphs[idx]
        text = para.text.strip()

        pPr = para._element.find(qn('w:pPr'))
        is_bullet = False
        if pPr is not None:
            is_bullet = pPr.find(qn('w:numPr')) is not None
            if not is_bullet:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    is_bullet = pStyle.get(qn('w:val')) == 'ListBullet'

        if not text:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            if not last_was_empty and html_parts:
                html_parts.append('<p style="margin: 0; line-height: 0.8;">&nbsp;</p>')
                last_was_empty = True
            continue

        last_was_empty = False

        run_html = ""
        for child in para._element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'r':
                t = ""
                for t_elem in child.findall(qn('w:t')):
                    t += t_elem.text or ""
                if not t:
                    continue
                t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                is_bold = False
                is_italic = False
                is_underline = False
                for run in para.runs:
                    if run._element is child:
                        is_bold = _is_run_bold(run)
                        is_italic = _is_run_italic(run)
                        is_underline = _is_run_underline(run)
                        break

                if is_bold:
                    t = f"<strong>{t}</strong>"
                if is_italic:
                    t = f"<em>{t}</em>"
                if is_underline:
                    t = f"<u>{t}</u>"
                run_html += t

            elif tag == 'hyperlink':
                rid = child.get(qn('r:id'))
                rel = doc.part.rels.get(rid) if rid else None
                url = ""
                if rel and hasattr(rel, 'target_ref'):
                    url = rel.target_ref
                link_text = ""
                for r in child.findall(qn('w:r')):
                    for t_elem in r.findall(qn('w:t')):
                        link_text += t_elem.text or ""
                if not link_text:
                    link_text = url
                link_text = link_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if url and url.startswith("http"):
                    run_html += f'<a href="{url}" target="_blank" style="color: #882345; text-decoration: underline;">{link_text}</a>'
                else:
                    run_html += link_text

        url_pattern = r'(?<!href=")(https?://[^\s<>"]+)'
        run_html = re.sub(url_pattern, r'<a href="\1" target="_blank" style="color: #882345;">\1</a>', run_html)

        if is_bullet:
            if not in_list:
                html_parts.append('<ul style="margin: 10px 0 10px 40px; padding: 0; list-style-type: disc;">')
                in_list = True
            run_html = run_html.replace("\n", "<br/>")
            html_parts.append(f'<li style="margin: 6px 0; padding-left: 4px; line-height: 1.6;">{run_html}</li>')
        else:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f'<p style="margin: 2px 0; line-height: 1.6;">{run_html}</p>')

    if in_list:
        html_parts.append("</ul>")

    return "".join(html_parts)

# ═══════════════════════════════════════════════════════════════
# FASTAPI APP
# ═══════════════════════════════════════════════════════════════

app = FastAPI(title="BC Admissions Email Assistant")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


class EmailRequest(BaseModel):
    email_text: str


class FeedbackRequest(BaseModel):
    query_id: str
    feedback: str


@app.get("/api/health")
def health():
    return {"status": "healthy", "templates": len(SECTIONS), "mongodb": db is not None}


@app.get("/api/templates")
def templates():
    return {
        "templates": [{"title": s["title"], "text": s["text"]} for s in SECTIONS],
        "count": len(SECTIONS),
    }


@app.post("/api/generate")
def generate(req: EmailRequest):
    email_text = req.email_text.strip()
    if len(email_text) < 3:
        raise HTTPException(400, "Email too short")

    student_info = extract_student_info(email_text)
    chosen = choose_template(email_text)

    if chosen is None:
        query_id = log_query({
            "topic": student_info["topic"],
            "template_title": "",
            "confidence": 0,
            "matched": False,
        })
        return {
            "success": False,
            "query_id": query_id,
            "student_name": student_info["name"] or "(not found)",
            "student_topic": student_info["topic"],
            "message": "No matching template found.",
        }

    working_doc = copy.deepcopy(ORIGINAL_DOC)
    working_section = {
        "title": chosen["title"],
        "start": chosen["start"],
        "end": chosen["end"],
        "para_indices": chosen["para_indices"],
    }
    fill_placeholders(working_doc, working_section, student_info)

    file_id = str(uuid.uuid4())[:8]
    docx_filename = f"response_{file_id}.docx"
    docx_path = OUTPUT_DIR / docx_filename
    export_doc(working_doc, working_section, str(docx_path))

    response_text = get_plain_text(working_doc, working_section)
    response_html = get_html_text(working_doc, working_section)

    old_files = sorted(OUTPUT_DIR.glob("response_*.docx"), key=lambda f: f.stat().st_mtime)
    for f in old_files[:-5]:
        f.unlink()

    query_id = log_query({
        "topic": student_info["topic"],
        "template_title": chosen["title"],
        "confidence": chosen.get("llm_confidence", 0),
        "matched": True,
    })

    return {
        "success": True,
        "query_id": query_id,
        "student_name": student_info["name"] or "(not found)",
        "student_semester": student_info["semester"] or "(not specified)",
        "student_topic": student_info["topic"],
        "template_title": chosen["title"],
        "response_text": response_text,
        "response_html": response_html,
        "docx_download_url": f"/api/download/{docx_filename}",
        "confidence": chosen.get("llm_confidence", 0),
        "message": "OK",
    }


@app.post("/api/feedback")
def submit_feedback(req: FeedbackRequest):
    if db is None:
        return {"success": False, "message": "Database not connected"}
    try:
        from bson import ObjectId
        db.queries.update_one(
            {"_id": ObjectId(req.query_id)},
            {"$set": {"feedback": req.feedback}}
        )
        return {"success": True}
    except Exception as e:
        return {"success": False, "message": str(e)}


@app.get("/api/dashboard")
def dashboard():
    if db is None:
        return {"success": False, "message": "Database not connected"}
    try:
        total = db.queries.count_documents({})
        matched = db.queries.count_documents({"matched": True})
        unmatched = db.queries.count_documents({"matched": False})
        thumbs_up = db.queries.count_documents({"feedback": "up"})
        thumbs_down = db.queries.count_documents({"feedback": "down"})

        pipeline = [
            {"$match": {"matched": True}},
            {"$group": {"_id": "$template_title", "count": {"$sum": 1}}},
            {"$sort": {"count": -1}},
            {"$limit": 10},
        ]
        top_templates = list(db.queries.aggregate(pipeline))

        recent = list(
            db.queries.find(
                {},
                {"_id": 0, "topic": 1, "template_title": 1, "confidence": 1, "matched": 1, "feedback": 1, "timestamp": 1}
            ).sort("timestamp", -1).limit(20)
        )
        for r in recent:
            if "timestamp" in r and r["timestamp"]:
                r["timestamp"] = r["timestamp"].isoformat() + "Z"

        avg_pipeline = [
            {"$match": {"matched": True}},
            {"$group": {"_id": None, "avg": {"$avg": "$confidence"}}},
        ]
        avg_result = list(db.queries.aggregate(avg_pipeline))
        avg_confidence = round(avg_result[0]["avg"], 1) if avg_result else 0

        return {
            "success": True,
            "total_queries": total,
            "matched": matched,
            "unmatched": unmatched,
            "thumbs_up": thumbs_up,
            "thumbs_down": thumbs_down,
            "avg_confidence": avg_confidence,
            "top_templates": [{"name": t["_id"], "count": t["count"]} for t in top_templates],
            "recent_queries": recent,
        }
    except Exception as e:
        return {"success": False, "message": str(e)}


@app.get("/api/download/{filename}")
def download(filename: str):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404, "File not found")
    return FileResponse(str(path), filename=filename)